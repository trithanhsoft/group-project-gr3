from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "PCS_Test_Plan_ISTQB_v1.2.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "2E75B6"
LIGHT = "D9EAF7"
PALE = "EEF4F8"
GRAY = "666666"
WHITE = "FFFFFF"
RED = "9C0006"
GOLD = "BF8F00"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.49)

def font(run, name="Arial", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for nm, sz, before, after, color in [("Heading 1",16,14,7,NAVY),("Heading 2",13,10,5,BLUE),("Heading 3",11.5,8,4,NAVY)]:
    s=styles[nm]; s.font.name="Arial"; s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn("w:ascii"),"Arial"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Arial")
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

for name in ["List Bullet", "List Number"]:
    s=styles[name]; s.font.name="Arial"; s.font.size=Pt(10.5); s.paragraph_format.left_indent=Inches(.5); s.paragraph_format.first_line_indent=Inches(-.25); s.paragraph_format.space_after=Pt(4)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def margins(cell, top=90, start=110, bottom=90, end=110):
    tc=cell._tc.get_or_add_tcPr(); mar=tc.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); tc.append(mar)
    for tag,val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        node=mar.find(qn("w:"+tag))
        if node is None: node=OxmlElement("w:"+tag); mar.append(node)
        node.set(qn("w:w"),str(val)); node.set(qn("w:type"),"dxa")

def set_cell_text(cell, text, bold=False, color=None, size=9.3, align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
    font(p.add_run(str(text)), size=size, bold=bold, color=color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.style="Table Grid"
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i],h,True,WHITE,9.2,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[i],NAVY)
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement("w:tblHeader"); repeat.set(qn("w:val"),"true"); trPr.append(repeat)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        cant=OxmlElement("w:cantSplit"); t.rows[-1]._tr.get_or_add_trPr().append(cant)
        for i,v in enumerate(row):
            align=WD_ALIGN_PARAGRAPH.CENTER if len(str(v))<18 and i>0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i],v,False,None,9.1,align)
            if ri%2: shade(cells[i],PALE)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def p(text="", bold=False, italic=False, color=None, size=10.5, align=None, after=6):
    x=doc.add_paragraph(); x.paragraph_format.space_after=Pt(after)
    if align is not None: x.alignment=align
    font(x.add_run(text),size=size,bold=bold,color=color,italic=italic)
    return x

def bullets(items):
    for x in items:
        q=doc.add_paragraph(style="List Bullet"); font(q.add_run(x),size=10.5)

def callout(label, text, fill=LIGHT, color=NAVY):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); shade(c,fill); margins(c,140,160,140,160); c.text=""
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0)
    font(q.add_run(label+": "),size=10,bold=True,color=color); font(q.add_run(text),size=10,color=color)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    font(paragraph.add_run("PCS Test Plan  |  "),size=8.5,color=GRAY)
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); paragraph._p.append(fld)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("PICKLEBALL COURT & COACH BOOKING SYSTEM"),size=8,bold=True,color=GRAY)
add_page_number(sec.footer.paragraphs[0])

# Cover
p("QUALITY ASSURANCE / ISTQB v4.0",True,False,BLUE,10,WD_ALIGN_PARAGRAPH.CENTER,24)
p("TEST PLAN",True,False,NAVY,28,WD_ALIGN_PARAGRAPH.CENTER,5)
p("API & AUTOMATION SCOPE",True,False,BLUE,16,WD_ALIGN_PARAGRAPH.CENTER,18)
p("Pickleball Court & Coach Booking System (PCS)",False,False,GRAY,13,WD_ALIGN_PARAGRAPH.CENTER,36)
table(["Document Control","Value"],[
    ["Version","v1.2"],["Prepared by","LÊ THỊ VĂN ANH - QA Leader"],["Reviewed by","TRƯƠNG QUANG TUÂN - Tester 1"],
    ["Approved by","TRẦN QUỐC SANG - Project Manager"],["Date","19/07/2026"],["Status","Review - current execution blockers recorded"]
],[1.8,4.7])
callout("Current verification", "35 tests executed and passed; 7 suites failed during collection because required packages were unavailable from the root test workspace. Exit criteria are therefore not yet met.", "FFF2CC", GOLD)
p("Internal project document",False,True,GRAY,9,WD_ALIGN_PARAGRAPH.CENTER,0)
doc.add_page_break()

doc.add_heading("Document information",0)
table(["Field","Value"],[
    ["Project Name","Pickleball Court & Coach Booking System (PCS)"],["Document Title","TEST PLAN - API & AUTOMATION SCOPE"],
    ["Version","v1.2"],["Prepared By","LÊ THỊ VĂN ANH - QA Leader"],["Reviewed By","TRƯƠNG QUANG TUÂN - Tester 1"],
    ["Approved By","TRẦN QUỐC SANG - Project Manager"],["Date","19/07/2026"],["Status","Review"]
],[1.8,4.7])
doc.add_heading("Revision history",1)
table(["Version","Date","Author","Description"],[
    ["1.0","10/10/2023","LÊ THỊ VĂN ANH","Initial API and automation test plan."],
    ["1.1","19/07/2026","LÊ THỊ VĂN ANH","Updated Vitest workspace and QA Dashboard scope."],
    ["1.2","19/07/2026","LÊ THỊ VĂN ANH","Reconciled technology versions and recorded current execution blockers."]
],[.7,1.05,1.65,3.1])
doc.add_page_break()
doc.add_heading("Contents",1)
for x in ["1. Introduction","2. Project Overview","3. Test Objectives","4. Test Scope","5. Test Items","6. Test Types","7. Test Approach / Strategy","8. Entry Criteria","9. Exit Criteria","10. Test Environment","11. Test Tools","12. Test Data Management","13. Roles and Responsibilities","14. Test Schedule","15. Defect Management Process","16. Risks and Mitigation","17. Test Deliverables","18. Communication Plan","19. Approvals","Appendices A-D"]:
    p(x,False,False,None,10,after=2)
doc.add_page_break()

doc.add_heading("1. INTRODUCTION",0)
p("This Test Plan defines the approach, scope, schedule, resources, controls, and completion criteria for API integration and automation testing of PCS. It provides a shared basis for planning, execution, defect triage, reporting, and release decisions.")
p("Testing protects the core business flow - account access, court and coach booking, promotion validation, payment confirmation, and operational reporting - while reducing regression risk across Next.js route handlers, service modules, SQL Server interactions, and external integrations.")
callout("Objective", "Provide repeatable, risk-based evidence that critical PCS workflows behave correctly and that known blockers are visible before release approval.")

doc.add_heading("2. PROJECT OVERVIEW",0)
p("PCS supports online court and coach reservations, player matching, payment processing, promotions, reviews, notifications, administrative reporting, and AI-assisted customer interaction. The solution uses separate frontend and backend Next.js applications plus a FastAPI-based AI service and a React/Vite QA dashboard.")
table(["Area","Verified repository baseline"],[
    ["Frontend","Next.js 15.x, React 18.3, TypeScript"],["Backend","Next.js 16.2, React 19.2, Node.js route handlers"],
    ["Database","Microsoft SQL Server through mssql; repositories mocked in automated tests"],["Authentication","JWT, bcrypt/bcryptjs, Google OAuth integration"],
    ["Payments","PayOS SDK/webhook routes; refund helper integrations"],["AI service","FastAPI/Python service for chatbot and player/coach scoring"],
    ["QA dashboard","Vite/React dashboard generated from test and coverage artifacts"]
],[1.35,5.15])

doc.add_heading("3. TEST OBJECTIVES",0)
bullets(["Verify critical authentication, court, booking, coach, payment, promotion, review, notification, reporting, matching, and AI rules.",
         "Validate API status codes, authorization gates, input validation, repository interactions, and external callback handling.",
         "Detect overlap, boundary, timezone, voucher, refund, and security defects early.",
         "Maintain a repeatable regression suite and publish evidence through JSON/HTML coverage and the PCS QA Dashboard.",
         "Target at least 90% statement coverage and 85% branch coverage for explicitly targeted business-logic modules."])

doc.add_heading("4. TEST SCOPE",0)
doc.add_heading("4.1 In scope",1)
table(["Module","Coverage focus"],[
    ["Authentication","Registration/login validation, duplicate checks, JWT issuance and access control"],
    ["Court & Booking","Court availability, time-slot overlap, daily limits, booking lifecycle and holds"],
    ["Coach","Coach profiles, approval status and appointment slots"],
    ["Payment & Refund","Checkout creation, PayOS webhook handling, status changes and refund rules"],
    ["Promotion","Active dates, minimum order thresholds and invalid voucher handling"],
    ["Other services","Reviews, notifications, reports, player matching and AI fallbacks"],
    ["Automation assets","Vitest workspace/configuration, shared test data, mocks, coverage output and dashboard scripts"]
],[1.45,5.05])
doc.add_heading("4.2 Out of scope",1)
bullets(["High-volume load, stress and endurance testing.","Production payment execution and production customer data.","Comprehensive mobile-device and cross-browser manual certification.","Infrastructure penetration testing and third-party vendor certification.","Formal UAT execution by customer stakeholders."])

doc.add_heading("5. TEST ITEMS",0)
table(["Item","Repository location / description"],[
    ["API test suites","tests/api/auth.api.test.ts, booking.api.test.ts, court.api.test.ts, payment.api.test.ts"],
    ["Unit suites","tests/unit/*.test.ts covering 12 service-oriented areas"],
    ["UI suite","frontend/tests/ui/login.ui.test.tsx"],
    ["Shared data and mocks","tests/data/testData.ts; tests/mock/*.ts; setupBackendTests.ts; setupTests.ts"],
    ["Payment route","backend/src/app/api/payments/payos-webhook/route.ts"],
    ["Workspace configuration","vitest.workspace.ts and vitest.config.ts"],
    ["Dashboard","scripts/generate-dashboard-data.ts, scripts/serve-dashboard.js, test-dashboard-src/"]
],[1.55,4.95])

doc.add_heading("6. TEST TYPES",0)
table(["Test type","Purpose"],[
    ["Unit testing","Isolate service/business rules using mocked repositories and dependencies."],
    ["API integration testing","Exercise route-handler behavior with request and repository mocks."],
    ["UI component testing","Validate login form behavior in JSDOM with React Testing Library."],
    ["Regression testing","Re-run automated suites after changes to detect unintended impact."],
    ["Security-focused testing","Check authentication, authorization, validation and webhook trust boundaries."],
    ["Coverage analysis","Measure statement, branch, function and line execution with V8 coverage."]
],[1.65,4.85])

doc.add_heading("7. TEST APPROACH / TEST STRATEGY",0)
p("Testing is automated-first, risk-based, and aligned with the repository's Vitest workspace. Backend suites run in a Node environment; frontend component tests run in JSDOM. Database and third-party behavior is isolated with mocks so that failures can be reproduced deterministically.")
table(["Level","Approach","Primary evidence"],[
    ["Unit","White-box branch and error-path validation","Vitest unit suite results"],
    ["API integration","Request/response and service interaction checks","API suite results and status assertions"],
    ["UI component","Black-box validation of login interactions","React Testing Library assertions"],
    ["Regression","Full workspace execution after material changes","JSON test report and dashboard"],
    ["Coverage","V8 instrumentation with excluded build/test artifacts","HTML/JSON coverage reports"]
],[1.15,3.15,2.2])
doc.add_heading("Test design techniques",1)
bullets(["Equivalence partitioning for valid/invalid credentials and promotion eligibility.","Boundary value analysis for time slots, daily booking limits, rating ranges and discount thresholds.","Decision tables for compound registration validation and payment state rules.","State transition testing for booking/payment lifecycle changes.","Use-case testing for book-court-and-pay end-to-end behavior."])

doc.add_heading("8. ENTRY CRITERIA",0)
bullets(["Requirements and acceptance criteria for the target module are available and reviewed.","Target route/service implementation compiles and is reachable by the test workspace.","Root and application dependencies required by imported modules are installed.","Shared mocks and deterministic test data are prepared.","Test environment variables contain only sandbox/non-production credentials.","The build under test is identified and deployed to the intended test environment."])

doc.add_heading("9. EXIT CRITERIA",0)
table(["Criterion","Target","Current assessment (19/07/2026)"],[
    ["Planned tests executed","100%","Not met - 7 suites failed during collection"],
    ["Executed tests pass","100%","Met for executed set - 35/35 passed"],
    ["Critical/high defects","0 open","Pending triage of collection blockers"],
    ["Targeted statement coverage",">= 90%","Baseline report states 92.50%; rerun required after blockers"],
    ["Targeted branch coverage",">= 85%","Baseline report states 88.75%; rerun required after blockers"],
    ["Dashboard and reports","Generated and reviewable","Artifacts exist; refresh after successful full run"]
],[1.65,1.25,3.6])
callout("Release gate", "The document remains in Review status until all suites collect and execute successfully, results are refreshed, and blocking defects are closed or formally accepted.", "FCE4D6", RED)

doc.add_heading("10. TEST ENVIRONMENT",0)
table(["Component","Configuration"],[
    ["Operating system","Windows 10/11 development environment"],["Runtime","Node.js; Python/FastAPI for AI service"],
    ["Test framework","Vitest 3.0.7 declared at repository root"],["Backend test environment","Node with backend aliases and setupBackendTests.ts"],
    ["Frontend test environment","JSDOM with setupTests.ts"],["Coverage provider","@vitest/coverage-v8 3.0.7"],
    ["Database","Mocked Microsoft SQL Server connections"],["External services","Mocked PayOS, email/OAuth, refund and AI dependencies"]
],[1.75,4.75])

doc.add_heading("11. TEST TOOLS",0)
table(["Tool","Purpose"],[
    ["Vitest","Unit, API integration, UI and regression execution"],["React Testing Library","DOM rendering and interaction assertions"],
    ["V8 coverage","Statement, branch, function and line metrics"],["TypeScript / tsx","Typed test and dashboard scripts"],
    ["Postman assets","Manual API collection/environment support"],["Vite + React dashboard","Visual test, traceability, coverage and defect reporting"],
    ["Git / npm scripts","Version control and repeatable execution commands"]
],[1.7,4.8])

doc.add_heading("12. TEST DATA MANAGEMENT",0)
p("Reusable payloads are centralized in tests/data/testData.ts and module-specific mocks in tests/mock/. Tests must reset mock state between cases, avoid real customer data, and use deterministic timestamps and identifiers where business rules are time-sensitive.")
table(["Dataset","Purpose","Control"],[
    ["Valid/invalid auth","Token and validation paths","Synthetic identities only"],
    ["Overlapping slots","Booking-conflict detection","Fixed court/date/time records"],
    ["Voucher boundaries","Expiry and order thresholds","Dates relative to controlled clock"],
    ["Payment callbacks","Signature and status transitions","Sandbox/mock payloads"],
    ["Large/edge records","Boundary and resilience paths","Generated, non-production data"]
],[1.45,2.55,2.5])

doc.add_heading("13. ROLES AND RESPONSIBILITIES",0)
table(["Role / person","Responsibilities"],[
    ["Project Manager - Trần Quốc Sang","Approve scope and deliverables; coordinate use cases and decision-table requirements."],
    ["QA Leader - Lê Thị Văn Anh","Own plan, framework, API automation, execution evidence, dashboard and defect lifecycle."],
    ["Tester 1 - Trương Quang Tuân","Develop unit tests, maintain workspace configuration and peer-review QA outputs."],
    ["Developer 1 - Nguyễn Đào Văn Quý","Implement/fix backend APIs, repository hooks and route defects."],
    ["Developer 2 - Lê Hữu Sơn","Implement/fix frontend login and validation behavior."]
],[2.35,4.15])

doc.add_heading("14. TEST SCHEDULE",0)
table(["Phase","Start","End","Owner"],[
    ["Test planning","01/10/2023","07/10/2023","Project Manager / QA Lead"],
    ["Vitest workspace setup","08/10/2023","12/10/2023","QA Lead / Tester 1"],
    ["API automation","13/10/2023","20/10/2023","QA Lead"],
    ["Coverage instrumentation","21/10/2023","24/10/2023","QA Lead"],
    ["Dashboard deployment","25/10/2023","27/10/2023","QA Lead"],
    ["Regression and defect review","28/10/2023","01/11/2023","QA Team"],
    ["v1.2 verification and closure","19/07/2026","TBD","QA Lead / Developers"]
],[2.25,1.0,1.0,2.25])

doc.add_heading("15. DEFECT MANAGEMENT PROCESS",0)
callout("Workflow", "New -> Assigned -> In Progress -> Fixed -> Retest -> Closed")
table(["Severity","Definition","Response expectation"],[
    ["Critical","Crash, data corruption, unauthorized access, or total loss of a core workflow.","Immediate triage; release blocker"],
    ["High","Payment/booking/security rule failure with no acceptable workaround.","Fix before release"],
    ["Medium","Incorrect non-critical behavior or reporting with a workaround.","Schedule and retest"],
    ["Low","Cosmetic, warning, wording or minor usability issue.","Backlog unless release-critical"]
],[1.0,3.45,2.05])
p("Each defect record must include an identifier, environment/build, summary, preconditions, reproducible steps, expected and actual results, severity, priority, owner, evidence, status, and retest result.")

doc.add_heading("16. RISKS AND MITIGATION",0)
table(["Risk","Impact","Mitigation / contingency"],[
    ["Root dependency resolution mismatch","High","Install or expose application dependencies to the root Vitest workspace; lock versions; rerun all suites."],
    ["Deprecated workspace configuration","Medium","Migrate vitest.workspace.ts projects into test.projects in vitest.config.ts before the next major Vitest upgrade."],
    ["Timezone drift","High","Pin Asia/Ho_Chi_Minh or use controlled clocks for time-sensitive cases."],
    ["Mock divergence from production","High","Review mock contracts against repository and provider interfaces each sprint; retain a small sandbox integration set."],
    ["Incomplete API routes","Medium","Mock stable interfaces to author tests early; track blocked cases explicitly."],
    ["Misleading aggregate coverage","Medium","Report targeted-module coverage separately from full-workspace coverage and document exclusions."]
],[1.65,.75,4.1])

doc.add_heading("17. TEST DELIVERABLES",0)
bullets(["Approved Test Plan (this document).","Automated unit, API integration, and UI test suites.","JSON/HTML coverage reports and machine-readable test results.","PCS QA Dashboard and generated dashboard data.","Test execution, summary, automation, data, traceability, decision-table and use-case reports.","Defect report with retest evidence.","Postman collection and environment for supported manual checks."])

doc.add_heading("18. COMMUNICATION PLAN",0)
table(["Activity","Frequency","Participants","Output"],[
    ["QA stand-up","Daily during execution","QA Lead, Tester 1","Progress, blockers, next actions"],
    ["Defect triage","Twice weekly / as needed","QA, developers, PM","Severity, ownership, target fix"],
    ["Status report","Weekly","QA Lead, PM, stakeholders","Execution, coverage, defects, risk"],
    ["Release review","At exit gate","PM, QA Lead, relevant owners","Go/no-go recommendation"]
],[1.35,1.3,1.9,1.95])

doc.add_page_break()
doc.add_heading("19. APPROVALS",0)
table(["Name","Role","Signature","Date"],[
    ["TRẦN QUỐC SANG","Project Manager","Pending v1.2 approval",""],
    ["LÊ THỊ VĂN ANH","QA Leader","Prepared","19/07/2026"],
    ["TRƯƠNG QUANG TUÂN","Tester 1 / Reviewer","Pending review",""]
],[1.7,1.7,1.8,1.3])

doc.add_page_break()
doc.add_heading("APPENDIX A - TEST CASE TEMPLATE",0)
table(["Field","Required content"],[
    ["Test Case ID","Unique identifier, e.g. TC_API_09"],["Module","Functional area under test"],
    ["Preconditions","Required initial state and data"],["Test Steps","Numbered execution actions"],
    ["Expected Result","Observable acceptance condition"],["Actual Result","Observed result and evidence"],
    ["Status","Pass / Fail / Blocked / Not Run"]
],[1.75,4.75])

doc.add_heading("APPENDIX B - DEFECT REPORT TEMPLATE",0)
table(["Field","Required content"],[
    ["Defect ID","Unique identifier, e.g. DF_PAY_01"],["Summary","Concise defect statement"],
    ["Steps to Reproduce","Repeatable sequence including data"],["Expected / Actual","Acceptance condition and observed behavior"],
    ["Severity / Priority","Critical-High-Medium-Low / P1-P2-P3"],["Status / Owner","Lifecycle state and accountable person"],
    ["Evidence","Logs, screenshots, response payloads or test output"]
],[1.75,4.75])

doc.add_heading("APPENDIX C - DECISION TABLE SAMPLE",0)
p("Registration validation example")
table(["Conditions / result","TC1","TC2","TC3","TC4","TC5","TC6","TC7","TC8"],[
    ["Email valid & unique","Y","Y","Y","Y","N","N","N","N"],
    ["Phone valid & unique","Y","Y","N","N","Y","Y","N","N"],
    ["Password strong","Y","N","Y","N","Y","N","Y","N"],
    ["Expected result","Success","Err Pass","Err Phone","Err Both","Err Email","Err Both","Err Both","Err All"]
],[2.1,.55,.55,.55,.55,.55,.55,.55,.55])

doc.add_heading("APPENDIX D - USE CASE TESTING SAMPLE",0)
p("UC-04: Book court and pay online")
table(["Step","User / system action","Expected result"],[
    ["1","Select court and 09:00-10:00 slot","Slot is validated and held according to booking rules."],
    ["2","Proceed to online payment","Backend creates a PayOS payment link/QR for the pending order."],
    ["3","PayOS sends signed callback","Webhook validates the callback before accepting the update."],
    ["4","Process valid payment result","Order becomes Paid and the reserved slot is unavailable to others."]
],[.65,2.85,3.0])

doc.add_heading("APPENDIX E - CURRENT VERIFICATION RECORD",0)
table(["Metric","Result"],[
    ["Command","npm test -- --reporter=json --outputFile=tmp-test-results.json"],
    ["Executed tests","35 passed, 0 failed, 0 pending"],["Suite collection","37 passed; 7 failed; overall run unsuccessful"],
    ["Collection blockers","Missing root-resolvable packages: @react-oauth/google, nodemailer, axios, exceljs, google-auth-library"],
    ["Configuration warning","vitest.workspace.ts is deprecated; migrate to test.projects"],
    ["Disposition","Resolve dependencies/configuration, rerun full suite, refresh coverage and seek approval"]
],[1.65,4.85])

# Keep headings with following content and prevent isolated short table rows where possible.
for para in doc.paragraphs:
    if para.style and para.style.name.startswith("Heading"):
        para.paragraph_format.keep_with_next=True

doc.core_properties.title="PCS Test Plan - API & Automation Scope"
doc.core_properties.subject="ISTQB v4.0-aligned test plan"
doc.core_properties.author="Lê Thị Văn Anh"
doc.core_properties.keywords="PCS, Test Plan, ISTQB, Vitest, API, Automation"
doc.save(OUT)
print(OUT)
